#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能日记本 - 浏览器版
使用 Flask 提供本地浏览器界面，样式接近微软官网风格。
"""

import os
import json
import string
import webbrowser
from datetime import datetime

from docx import Document
from flask import Flask, abort, jsonify, render_template, request, send_from_directory

from weather_service import WeatherService
from location_service import LocationService, CHINESE_CITIES
from news_service import NewsService

app = Flask(__name__, template_folder="templates", static_folder="static")
weather_service = WeatherService()
location_service = LocationService()
news_service = NewsService()

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
SAVE_DIR = os.path.join(BASE_DIR, "saved_diaries")
SAVE_OPTIONS = [
    "saved_diaries",
    "saved_diaries/archives",
    "saved_diaries/drafts",
]
DEFAULT_SAVE_DIR = SAVE_OPTIONS[0]
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
os.makedirs(SAVE_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)
for option in SAVE_OPTIONS:
    os.makedirs(os.path.join(BASE_DIR, option), exist_ok=True)


def safe_filename(value):
    value = str(value or "").strip()
    allowed = "-_.() %s%s" % (string.ascii_letters, string.digits)
    return "".join(c if c in allowed else "_" for c in value)[:120]


def format_txt(payload):
    lines = [
        "智能日记本 - 浏览器版",
        "=" * 56,
        f"标题: {payload.get('title', '日记')}",
        f"城市: {payload.get('city', '未知')}",
        f"保存时间: {payload.get('saved_at', '')}",
        "",
        "天气信息:",
        f"  状态: {payload.get('weather', {}).get('condition', '未知')}",
        f"  温度: {payload.get('weather', {}).get('temperature_c', '--')}°C",
        f"  体感: {payload.get('weather', {}).get('feelslike_c', '--')}°C",
        f"  湿度: {payload.get('weather', {}).get('humidity', '--')}%",
        f"  风速: {payload.get('weather', {}).get('wind_kph', '--')} km/h",
        "",
        "日记内容:",
        payload.get('content', ''),
        "",
    ]

    images = payload.get('images', []) or []
    videos = payload.get('videos', []) or []

    if images:
        lines.append("图片列表:")
        lines.extend([f"  - {item}" for item in images])
        lines.append("")

    if videos:
        lines.append("视频列表:")
        lines.extend([f"  - {item}" for item in videos])
        lines.append("")

    lines.append("=" * 56)
    return "\n".join(lines)


def resolve_save_dir(save_dir):
    target = str(save_dir or DEFAULT_SAVE_DIR).strip() or DEFAULT_SAVE_DIR
    normalized = os.path.normpath(target)
    if os.path.isabs(normalized) or normalized.startswith('..') or normalized.startswith('/') or normalized.startswith('\\'):
        return None

    save_path = os.path.abspath(os.path.join(BASE_DIR, normalized))
    if not save_path.startswith(BASE_DIR):
        return None
    return save_path


def save_docx(payload, path):
    document = Document()
    document.add_heading('智能日记本 - 浏览器版', level=1)

    document.add_paragraph(f"标题: {payload.get('title', '日记')}")
    document.add_paragraph(f"城市: {payload.get('city', '未知')}")
    document.add_paragraph(f"保存时间: {payload.get('saved_at', '')}")

    document.add_heading('天气信息', level=2)
    weather = payload.get('weather', {})
    document.add_paragraph(f"状态: {weather.get('condition', '未知')}")
    document.add_paragraph(f"温度: {weather.get('temperature_c', '--')}°C")
    document.add_paragraph(f"体感: {weather.get('feelslike_c', '--')}°C")
    document.add_paragraph(f"湿度: {weather.get('humidity', '--')}%")
    document.add_paragraph(f"风速: {weather.get('wind_kph', '--')} km/h")

    document.add_heading('日记内容', level=2)
    document.add_paragraph(payload.get('content', ''))

    images = payload.get('images', []) or []
    videos = payload.get('videos', []) or []
    if images:
        document.add_heading('图片列表', level=2)
        for item in images:
            document.add_paragraph(item, style='List Bullet')

    if videos:
        document.add_heading('视频列表', level=2)
        for item in videos:
            document.add_paragraph(item, style='List Bullet')

    document.save(path)


def save_diary(payload, save_dir=None):
    target_dir = resolve_save_dir(save_dir)
    if not target_dir:
        target_dir = SAVE_DIR

    os.makedirs(target_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    title = safe_filename(payload.get('title', '日记')) or 'diary'
    city = safe_filename(payload.get('city', '未知')) or 'unknown'
    base_name = f"{title}_{city}_{timestamp}"
    json_filename = f"{base_name}.json"
    txt_filename = f"{base_name}.txt"
    docx_filename = f"{base_name}.docx"

    json_path = os.path.join(target_dir, json_filename)
    txt_path = os.path.join(target_dir, txt_filename)
    docx_path = os.path.join(target_dir, docx_filename)

    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)

    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(format_txt(payload))

    save_docx(payload, docx_path)

    return json_filename, txt_filename, docx_filename


@app.route('/')
def index():
    default_city = location_service.get_city_by_ip() or '北京'
    return render_template(
        'index.html',
        default_city=default_city,
        cities=CHINESE_CITIES,
        save_dirs=SAVE_OPTIONS,
        default_save_dir=DEFAULT_SAVE_DIR,
        base_dir=BASE_DIR,
    )


@app.route('/api/weather')
def api_weather():
    city = request.args.get('city', '').strip() or '北京'
    weather = weather_service.get_current_weather(city)
    if not weather:
        return jsonify(success=False, error='无法获取天气信息，请检查网络'), 503
    return jsonify(success=True, weather=weather)


@app.route('/api/news')
def api_news():
    news = news_service.get_hot_news()
    return jsonify(success=True, **news)


@app.route('/api/geocode')
def api_geocode():
    lat = request.args.get('lat')
    lon = request.args.get('lon')
    try:
        latitude = float(lat)
        longitude = float(lon)
    except (TypeError, ValueError):
        return jsonify(success=False, error='无效坐标'), 400

    city = location_service.get_city_by_coordinates(latitude, longitude)
    if not city:
        return jsonify(success=False, error='无法解析城市'), 404

    return jsonify(success=True, city=city)


@app.route('/api/save', methods=['POST'])
def api_save():
    data = request.json or {}
    content = data.get('content', '').strip()
    if not content:
        return jsonify(success=False, error='请填写日记内容'), 400

    save_dir = data.get('save_dir', DEFAULT_SAVE_DIR)
    if not resolve_save_dir(save_dir):
        return jsonify(success=False, error='保存目录无效，请选择正确的目录'), 400

    payload = {
        'title': data.get('title', '日记').strip() or '日记',
        'city': data.get('city', '北京').strip() or '北京',
        'weather': data.get('weather', {}),
        'content': content,
        'images': data.get('images', []),
        'videos': data.get('videos', []),
        'saved_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
    }

    json_filename, txt_filename, docx_filename = save_diary(payload, save_dir=save_dir)
    return jsonify(
        success=True,
        json_file=json_filename,
        txt_file=txt_filename,
        docx_file=docx_filename,
        save_dir=save_dir,
        message='保存成功'
    )


@app.route('/api/upload_media', methods=['POST'])
def upload_media():
    media_type = request.form.get('type')
    if media_type not in ('image', 'video'):
        return jsonify(success=False, error='不支持的媒体类型'), 400

    files = request.files.getlist('files')
    if not files:
        return jsonify(success=False, error='未检测到上传文件'), 400

    saved_paths = []
    for uploaded in files:
        filename = safe_filename(uploaded.filename)
        if not filename:
            continue

        prefix = 'img' if media_type == 'image' else 'vid'
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S%f')
        saved_name = f"{prefix}_{timestamp}_{filename}"
        saved_path = os.path.join(UPLOAD_DIR, saved_name)
        uploaded.save(saved_path)
        saved_paths.append(f"uploads/{saved_name}")

    if not saved_paths:
        return jsonify(success=False, error='文件保存失败'), 500

    return jsonify(success=True, items=saved_paths)


@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(UPLOAD_DIR, filename)


@app.route('/download/<path:save_dir>/<path:filename>')
def download(save_dir, filename):
    save_path = resolve_save_dir(save_dir)
    if not save_path:
        abort(404)
    return send_from_directory(save_path, filename, as_attachment=True)


def open_browser(url):
    try:
        webbrowser.open(url, new=2, autoraise=True)
    except Exception:
        pass


def main():
    url = 'http://127.0.0.1:5000'
    open_browser(url)
    app.run(host='127.0.0.1', port=5000, debug=False)


if __name__ == '__main__':
    main()
